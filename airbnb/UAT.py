from docx import Document
from datetime import date

# Create a new Word document
doc = Document()
doc.add_heading('User Acceptance Test Document', 0)

# Project details
doc.add_paragraph('Project: Data Warehouse Validation for Airbnb Dataset')
doc.add_paragraph('Version: 1.0')
doc.add_paragraph(f'Date: {date.today().strftime("%B %d, %Y")}')
doc.add_paragraph('Prepared by: [Your Name]')

# Section 1: Purpose
doc.add_heading('1. Purpose', level=1)
doc.add_paragraph(
    "This document outlines the acceptance testing process to validate that the data pipeline correctly loads and "
    "transforms data from the Data Lake to the Data Warehouse. This ensures data integrity, accuracy, and completeness."
)

# Section 2: Scope
doc.add_heading('2. Scope', level=1)
doc.add_paragraph(
    "The scope of this test includes:\n"
    "- Verification that today's data in the Data Lake matches today's data in the Data Warehouse.\n"
    "- Batch ID consistency validation between the Data Lake and the Data Warehouse.\n"
    "- Basic row count checks for dimensional and fact tables."
)

# Section 3: Acceptance Criteria
doc.add_heading('3. Acceptance Criteria', level=1)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Test ID'
hdr_cells[1].text = 'Test Description'
hdr_cells[2].text = 'Pass Criteria'

criteria = [
    ('TC001', 'Row Count Check - Source vs Target', 'Row counts must match between source and target tables.'),
    ('TC002', 'Data Match Check - Source vs Target (by Surrogate Key)', 'No mismatches in key fields or row data.'),
    ('TC003', 'Batch ID Validation', 'Latest batch ID must match in both Data Lake and Warehouse.')
]

for test_id, desc, pass_criteria in criteria:
    row_cells = table.add_row().cells
    row_cells[0].text = test_id
    row_cells[1].text = desc
    row_cells[2].text = pass_criteria

# Section 4: Test Cases
doc.add_heading('4. Test Cases', level=1)

# TC001
doc.add_heading('Test Case TC001: Row Count Validation', level=2)
doc.add_paragraph(
    "Objective: Verify row count parity between source and target.\n"
    "Tables: RAW.REVIEWS → DEV.FCT_REVIEWS (and similar for others)\n"
    "Expected Result: Row counts match exactly."
)

# TC002
doc.add_heading('Test Case TC002: Data Match on Key and Row Hash', level=2)
doc.add_paragraph(
    "Objective: Verify data accuracy at row level using surrogate key comparison.\n"
    "Logic:\n"
    "- Generate row hashes based on key columns (REVIEW_SK or others).\n"
    "- Ensure each key exists in both tables.\n"
    "- Confirm that row hashes are identical for matching keys.\n"
    "Expected Result: Zero mismatches in keys and rows."
)

# TC003
doc.add_heading('Test Case TC003: Batch ID Match Validation', level=2)
doc.add_paragraph(
    "Objective: Confirm that the most recent batch ID loaded into the Data Warehouse matches the latest batch in the Data Lake.\n"
    "Validation: Query both layers for the max batch ID.\n"
    "Expected Result: MAX(batch_id) matches in both source and target."
)

# Section 5: Test Execution and Sign-off
doc.add_heading('5. Test Execution and Sign-off', level=1)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Test ID'
hdr_cells[1].text = 'Pass/Fail'
hdr_cells[2].text = 'Notes/Remarks'
hdr_cells[3].text = 'Tested By'
hdr_cells[4].text = 'Date'

for test_id in ['TC001', 'TC002', 'TC003']:
    row_cells = table.add_row().cells
    row_cells[0].text = test_id
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = ''
    row_cells[4].text = ''

# Section 6: Approval
doc.add_heading('6. Approval', level=1)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Role'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Signature'
hdr_cells[3].text = 'Date'

roles = ['Data Engineer', 'Data Owner / SME', 'Project Manager']

for role in roles:
    row_cells = table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = ''

# Save the document
file_path = 'UAT_Data_Validation.docx'
doc.save(file_path)
file_path
